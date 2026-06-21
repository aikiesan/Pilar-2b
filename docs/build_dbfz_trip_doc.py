#!/usr/bin/env python3
"""
Generator for the DBFZ Germany trip preparation dossier (.docx).

Audience: Lucas Nakamura Cerejo (NIPE-UNICAMP / CP2B), geospatial & data analyst,
lead developer of PILAR-2b, preparing a 2-day visit to DBFZ (Leipzig) and meetings
with Dr. Friederike Naegeli de Torres to prospect a future 8-12 month international
post-doctoral internship.

Run: python3 docs/build_dbfz_trip_doc.py
Output: docs/DBFZ_Germany_Trip_Preparation.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------------------------------------------------------------------
# Palette (DBFZ-ish green + a deep slate for headings)
# ---------------------------------------------------------------------------
GREEN = RGBColor(0x2E, 0x7D, 0x32)      # primary accent
DARK = RGBColor(0x1B, 0x2A, 0x3A)       # heading slate
GREY = RGBColor(0x55, 0x5F, 0x6B)       # secondary text
LIGHT_FILL = "EAF3EC"                    # table header fill
ALT_FILL = "F4F8F5"                      # zebra fill

doc = Document()

# Base style
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.12

# Heading styles
for name, size, color in [
    ("Heading 1", 17, GREEN),
    ("Heading 2", 13.5, DARK),
    ("Heading 3", 11.5, DARK),
]:
    st = doc.styles[name]
    st.font.name = "Calibri"
    st.font.size = Pt(size)
    st.font.color.rgb = color
    st.font.bold = True
    st.paragraph_format.space_before = Pt(12)
    st.paragraph_format.space_after = Pt(4)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------
def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), fill)
    tcPr.append(sh)


def set_cell_text(cell, text, bold=False, color=None, size=10, align=None, fill=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = color
    if fill:
        shade(cell, fill)
    return cell


def para(text="", bold=False, italic=False, color=None, size=None, align=None,
         space_after=None, space_before=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    if text:
        r = p.add_run(text)
        r.font.bold = bold
        r.font.italic = italic
        if color:
            r.font.color.rgb = color
        if size:
            r.font.size = Pt(size)
    return p


def rich(parts, style=None, align=None, space_after=None):
    """parts = list of (text, {bold,italic,color,size})"""
    p = doc.add_paragraph(style=style)
    if align:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    for text, fmt in parts:
        r = p.add_run(text)
        r.font.bold = fmt.get("bold", False)
        r.font.italic = fmt.get("italic", False)
        if fmt.get("color"):
            r.font.color.rgb = fmt["color"]
        if fmt.get("size"):
            r.font.size = Pt(fmt["size"])
    return p


def bullet(text, level=0, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    if bold_lead:
        r = p.add_run(bold_lead)
        r.font.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def numbered(text, bold_lead=None):
    p = doc.add_paragraph(style="List Number")
    if bold_lead:
        r = p.add_run(bold_lead)
        r.font.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def hrule():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E7D32")
    pbdr.append(bottom)
    pPr.append(pbdr)
    p.paragraph_format.space_after = Pt(2)
    return p


def make_table(headers, rows, widths=None, header_fill=LIGHT_FILL, zebra=True,
               font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, color=DARK, size=font_size,
                      fill=header_fill)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci, val in enumerate(row):
            set_cell_text(cells[ci], val, size=font_size)
            if zebra and ri % 2 == 1:
                shade(cells[ci], ALT_FILL)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    return t


def callout(title, body_lines, fill="EAF3EC", title_color=GREEN):
    """Single-cell shaded box used for tips / key messages."""
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    cell = t.rows[0].cells[0]
    shade(cell, fill)
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.font.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = title_color
    for line in body_lines:
        bp = cell.add_paragraph()
        bp.paragraph_format.space_before = Pt(2)
        rr = bp.add_run(line)
        rr.font.size = Pt(10)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def page_break():
    doc.add_page_break()


# ===========================================================================
# COVER
# ===========================================================================
sp = doc.add_paragraph()
sp.paragraph_format.space_before = Pt(60)

para("PREPARATION DOSSIER", bold=True, color=GREEN, size=12,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para("Research Visit & Collaboration Mission to the DBFZ — Germany",
     bold=True, color=DARK, size=24, align=WD_ALIGN_PARAGRAPH.CENTER,
     space_after=2)
para("Deutsches Biomasseforschungszentrum (German Biomass Research Centre), Leipzig",
     italic=True, color=GREY, size=12, align=WD_ALIGN_PARAGRAPH.CENTER,
     space_after=18)

para("Two days of meetings · facility tour · data-integration working sessions",
     color=DARK, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para("Prospecting a 2027 international post-doctoral internship (8–12 months)",
     color=GREEN, bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER,
     space_after=24)

# cover meta table
meta = doc.add_table(rows=0, cols=2)
meta.alignment = WD_TABLE_ALIGNMENT.CENTER
meta_rows = [
    ("Prepared for", "Lucas Nakamura Cerejo"),
    ("Role", "Geospatial & Data Analyst · Lead Developer, PILAR-2b"),
    ("Home institution", "NIPE — UNICAMP · CP2B (Campinas, SP, Brazil)"),
    ("Funding", "FAPESP 2024/01112-1 (CP2Bsd)"),
    ("Host (target)", "DBFZ — Bioenergy Systems Dept. · Working Group “Resources”"),
    ("Key counterpart", "Dr. Friederike Naegeli de Torres (Head, WG Resources)"),
    ("Document date", "21 June 2026"),
]
for k, v in meta_rows:
    cells = meta.add_row().cells
    set_cell_text(cells[0], k, bold=True, color=DARK, size=10, fill=LIGHT_FILL)
    set_cell_text(cells[1], v, size=10)
    cells[0].width = Inches(1.9)
    cells[1].width = Inches(4.3)

para("", space_after=18)
para("Confidential working document — for meeting preparation. Contains draft "
     "proposals and talking points to be refined before circulation.",
     italic=True, color=GREY, size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER)

page_break()

# ===========================================================================
# TABLE OF CONTENTS (static)
# ===========================================================================
para("Contents", bold=True, color=GREEN, size=16, space_after=6)
toc_items = [
    "1. Executive summary — the mission in one page",
    "2. Your objectives and how to frame them (the “ask”)",
    "3. Who you are meeting — DBFZ, the department, and Dr. Naegeli de Torres",
    "4. DBFZ expertise & developments you must know",
    "5. The bridge — PILAR-2b mapped onto DBFZ’s work",
    "6. Collaboration menu — concrete proposals, ranked",
    "7. The post-doctoral internship — routes, funding, and a draft project",
    "8. Proposed two-day agenda",
    "9. Data-integration working session — technical preparation",
    "10. Questions to ask them",
    "11. Questions they will ask you — and your answers",
    "12. What to prepare and bring (checklist)",
    "13. Logistics, language and cultural notes",
    "14. Follow-up plan (the 30 days after)",
    "Appendix A — One-page collaboration concept note (draft)",
    "Appendix B — Draft post-doc research outline",
    "Appendix C — Elevator pitch & key talking points",
    "Appendix D — Sources & useful links",
]
for it in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(it)
    r.font.size = Pt(10.5)
    r.font.color.rgb = DARK

page_break()

# ===========================================================================
# 1. EXECUTIVE SUMMARY
# ===========================================================================
doc.add_heading("1. Executive summary — the mission in one page", level=1)
hrule()

para("You are travelling to the DBFZ in Leipzig — Germany’s national biomass "
     "research centre (200+ staff, ~120 scientific projects/year) — for two days "
     "of meetings, a facility tour, and hands-on data-integration discussions. Your "
     "strategic goal is to convert an existing affinity into a formal collaboration "
     "and, from it, a funded 8–12 month post-doctoral internship at DBFZ in 2027.")

callout(
    "Why this fit is unusually strong",
    [
        "• Dr. Friederike Naegeli de Torres leads the DBFZ Working Group “Resources” — "
        "a geographer and GIS / remote-sensing specialist whose guiding question is "
        "literally: “Where can we find how much of which kind of biomass, how much is "
        "mobilisable, and what are the utilisation paths?” That is precisely what "
        "PILAR-2b answers for São Paulo.",
        "• DBFZ’s flagship public tool — the Resource Database (resDB) on the DBFZ "
        "DataLab — is the German national analogue of PILAR-2b. Your own README already "
        "cites it as a UX inspiration. You arrive as a peer, not a petitioner.",
        "• She has direct Brazil research experience (Atlantic Forest carbon, "
        "INTECRAL pasture-rehabilitation project, Rio de Janeiro) and DBFZ has a "
        "standing Brazil network (EMBRAPA, GIZ, Rio Grande do Sul universities, "
        "sugarcane/vinasse residues). You extend that network into São Paulo.",
    ],
)

para("The three things to land in two days", bold=True, color=DARK, size=11,
     space_before=8, space_after=2)
numbered("Establish personal and methodological rapport — show that PILAR-2b and "
         "the DBFZ resDB are two implementations of the same idea (theoretical → "
         "mobilisable biomass potential), and that your FDE correction-factor method "
         "is a contribution, not a duplication.", bold_lead="Rapport & credibility. ")
numbered("Agree on one or two concrete, low-friction first collaborations "
         "(a co-authored method paper, a joint conference session, an API/data "
         "interoperability pilot) that can start before any post-doc.",
         bold_lead="A first joint deliverable. ")
numbered("Secure a named host and an in-principle agreement to support a 2027 "
         "post-doc, plus clarity on the funding route (FAPESP BEPE, CAPES-Humboldt, "
         "DAAD, or DFG/BMBF bilateral).", bold_lead="A path to the post-doc. ")

para("Treat the post-doc as the natural consequence of a collaboration you start "
     "now — not as the opening request. Lead with mutual scientific value; the "
     "internship follows.", italic=True, color=GREY, space_before=6)

page_break()

# ===========================================================================
# 2. OBJECTIVES
# ===========================================================================
doc.add_heading("2. Your objectives and how to frame them (the “ask”)", level=1)
hrule()

doc.add_heading("2.1 Your layered objectives", level=2)
make_table(
    ["Horizon", "Objective", "What “success” looks like by end of visit"],
    [
        ["Immediate\n(these 2 days)",
         "Build rapport; demo PILAR-2b; map mutual interests; identify a first joint activity.",
         "A named DBFZ contact who is enthusiastic, plus one agreed next step in writing (even an email)."],
        ["Short term\n(2026)",
         "One tangible collaboration: co-authored methods note, joint conference session (FOSS4G Europe 2026), or data-interoperability pilot.",
         "A draft outline / shared doc and a target venue or deadline."],
        ["Medium term\n(2027)",
         "Funded 8–12 month post-doctoral internship hosted by WG Resources / Bioenergy Systems Dept.",
         "An in-principle yes from a host + an identified funding instrument + a timeline."],
        ["Long term",
         "Standing UNICAMP–DBFZ research bridge; PILAR-2b as the São Paulo / Brazil node of a resDB-style monitoring network.",
         "Shared roadmap and intent to pursue a joint bilateral grant."],
    ],
    widths=[1.2, 2.7, 2.6],
)

doc.add_heading("2.2 Your value proposition — what you bring to DBFZ", level=2)
para("Be explicit about this. You are not asking for charity; you bring assets DBFZ "
     "values and a geography they want to reach.")
bullet("a production-grade, open-source geospatial platform (PILAR-2b) already live "
       "and INPI-registered, covering 645 municipalities of São Paulo;",
       bold_lead="A working platform — ")
bullet("a feedstock-dependent efficiency (FDE) correction-factor methodology that "
       "decomposes theoretical potential into mobilisable potential — a methodological "
       "contribution that complements the DBFZ potential cascade;",
       bold_lead="A novel method — ")
bullet("hands-on full-stack engineering (FastAPI + PostGIS, Next.js/React-Leaflet, "
       "APIs, Sankey/IO economic modelling) — i.e. you can actually build and "
       "integrate, not only theorise;", bold_lead="Build capability — ")
bullet("access to curated Brazilian datasets (IBGE, MapBiomas, ANP/ANEEL, agricultural "
       "and livestock residues) and the São Paulo bioenergy context;",
       bold_lead="Data & context — ")
bullet("knowledge of spatial decision-support systems (SDSS) and public-policy "
       "application — the ‘so-what’ layer that turns potential maps into investment "
       "and policy decisions.", bold_lead="SDSS & policy framing — ")

callout(
    "Framing tip — the one sentence to rehearse",
    ["“PILAR-2b does for São Paulo’s residues what the DBFZ Resource Database does for "
     "Germany — and the FDE method I built to get from theoretical to mobilisable "
     "potential is exactly the kind of harmonisation a Brazil–Germany comparison needs. "
     "I’d like to explore doing that together.”"],
    fill="FFF4E5", title_color=RGBColor(0xB0, 0x6A, 0x00),
)

page_break()

# ===========================================================================
# 3. WHO YOU ARE MEETING
# ===========================================================================
doc.add_heading("3. Who you are meeting — DBFZ, the department, and Dr. Naegeli de Torres",
                level=1)
hrule()

doc.add_heading("3.1 The institution — DBFZ at a glance", level=2)
make_table(
    ["Attribute", "Detail"],
    [
        ["Full name", "Deutsches Biomasseforschungszentrum gemeinnützige GmbH (German Biomass Research Centre)"],
        ["Location", "Leipzig, Saxony, Germany (HQ at Torgauer Straße 116)"],
        ["Scale", "200+ employees; ~120 scientific bioenergy projects per year"],
        ["Mandate", "Technical, economic and ecological questions of biomass as an energy and material resource; advises the German federal government"],
        ["Core infrastructure", "Labs for biochemical conversion, thermo-chemical conversion, biorefineries; a Data Laboratory (DataLab) for data science / IT methods"],
        ["Public flagship tool", "DBFZ Resource Database (resDB) — national biomass monitoring, open web app + API"],
        ["Relevant department", "Bioenergy Systems Department — supply scenarios, system analysis, modelling of competing uses, biomass potential analysis"],
    ],
    widths=[1.7, 4.8],
)

doc.add_heading("3.2 The department & working group you are targeting", level=2)
rich([("Bioenergy Systems Department", {"bold": True}),
      (" — develops supply scenarios and utilisation strategies for bioenergy, "
       "models potential competing uses, and runs system analysis of demand-oriented "
       "bioenergy provision. This is the analytical, data- and model-heavy side of "
       "DBFZ — your natural home.", {})])
rich([("Working Group “Resources”", {"bold": True, "color": GREEN}),
      (" (led by Dr. Naegeli de Torres since 2023) — answers the resource-mapping "
       "question: where the biomass is, how much is mobilisable, and which utilisation "
       "paths exist. It is the group behind the spatial / potential-atlas side of the "
       "Resource Database. This is the single best-matched team for your profile.", {})])

doc.add_heading("3.3 Dr. Friederike Naegeli de Torres — your key counterpart", level=2)
make_table(
    ["", ""],
    [
        ["Role", "Head of the Working Group “Resources”, DBFZ (since 2023); researcher & GIS expert at DBFZ since 2020"],
        ["Background", "Geographer; MSc in Technology & Resource Management in the Tropics and Subtropics; PhD in Remote Sensing & Geoinformatics"],
        ["Core question", "“Where can we find how much of which kind of biomass? How much is mobilisable, and what are the utilisation paths?”"],
        ["International work", "Brazil, Colombia, Togo, Ethiopia — strong tropical/subtropical and Global-South focus"],
        ["Brazil research", "Carbon-storage capacity of a fragmented Atlantic-Forest landscape (Rio de Janeiro); pasture rehabilitation within the Brazilian–German INTECRAL project; pastures as climate-mitigation opportunity areas"],
        ["Contact", "+49 (0)341 2434-620 · DBFZ staff directory (Personalverzeichnis)"],
    ],
    widths=[1.5, 5.0],
)

callout(
    "Personal-rapport angles (use sparingly, sincerely)",
    ["• Shared methodological DNA — GIS + remote sensing applied to biomass/land "
     "resources. Speak her language: spatial potential assessment, mobilisable "
     "fraction, utilisation paths.",
     "• Genuine Brazil connection — she has done fieldwork in Brazil and works on "
     "tropics/subtropics; the surname “de Torres” and INTECRAL suggest comfort with "
     "the Brazilian research context. A few words in Portuguese or on the Atlantic "
     "Forest will land well.",
     "• Land-use & carbon — her pasture/carbon work connects to your land-use "
     "(MapBiomas) layer and to the sustainability framing of residue mobilisation."],
)

page_break()

# ===========================================================================
# 4. DBFZ EXPERTISE & DEVELOPMENTS
# ===========================================================================
doc.add_heading("4. DBFZ expertise & developments you must know", level=1)
hrule()
para("Know these well enough to reference them by name. They are the hooks your "
     "proposals hang on.")

doc.add_heading("4.1 The DBFZ Resource Database (resDB) — your anchor reference", level=2)
bullet("Germany’s national biomass monitoring, open to the public as a free web "
       "application on the DBFZ DataLab (datalab.dbfz.de/resdb).")
bullet("Covers ~77 biomass types across agriculture & forestry, municipal waste & "
       "sewage sludge, and industrial residues — with time series (2010–2020).")
bullet("Includes a biomass potential atlas tracing individual biomasses across the "
       "EU, plus documented calculation methodology and data basis per material.")
bullet("Exposes an API for integration into third-party applications — directly "
       "relevant to a PILAR-2b ↔ resDB interoperability pilot.")
bullet("Reports theoretical, technical and mobilisable potentials — the same cascade "
       "logic as your FDE method.")

doc.add_heading("4.2 Biomass potential methodology — the shared backbone", level=2)
para("DBFZ formalises the potential cascade and develops indicators of efficiency and "
     "sustainability (economic, ecological, technical). They also work on cascading "
     "use — ensuring material use is not blocked by energy use. Your FDE method "
     "(collection efficiency, competing uses, seasonal availability, logistics) is a "
     "transparent, factor-explicit way to perform exactly the theoretical→mobilisable "
     "step. Position FDE as harmonisable with — not competing against — their cascade.")

make_table(
    ["DBFZ potential cascade", "PILAR-2b / FDE equivalent"],
    [
        ["Theoretical potential", "Theoretical biomass potential (gross residue mass × BMP/VS/TS parameters)"],
        ["Technical potential", "Collection-efficiency & logistical-constraint factors of the FDE"],
        ["Mobilisable / sustainable potential", "Mobilisable potential after competing-uses & seasonal-availability factors"],
        ["Utilisation paths", "Technology-routes comparison (10+ biogas conversion pathways) + co-digestion C:N optimisation"],
        ["Regional value / system analysis", "IBGE 67-sector Leontief input-output model + payback/viability calculator"],
    ],
    widths=[3.0, 3.5],
)

doc.add_heading("4.3 DBFZ international & Brazil activities — you fit an existing strategy",
                level=2)
bullet("DBFZ runs cooperation networks across Latin America — especially Brazil, "
       "Chile and Mexico.")
bullet("In Brazil: a research & consultancy network with Rio Grande do Sul "
       "universities, EMBRAPA Suínos e Aves (Concórdia) and other EMBRAPA units, and "
       "GIZ Brazil — focused on biogas from agricultural residues and manure.")
bullet("Long-standing interest in sugarcane/ethanol residues (vinasse), manioc, "
       "manure and municipal waste — all squarely within São Paulo’s residue profile "
       "and PILAR-2b’s scope.")
bullet("Stated aim to support pilot plants and to grow the university network through "
       "continuous exchange of scientists and PhD students — i.e. they explicitly want "
       "people like you.")

callout(
    "Strategic read",
    ["DBFZ’s existing Brazil footprint is concentrated in the South (Rio Grande do "
     "Sul) and around sugarcane. PILAR-2b extends that footprint to São Paulo — the "
     "largest sugarcane/bioenergy state — with a ready-made open platform. You are a "
     "geographic and technological extension of a strategy they already fund."],
)

page_break()

# ===========================================================================
# 5. THE BRIDGE
# ===========================================================================
doc.add_heading("5. The bridge — PILAR-2b mapped onto DBFZ’s work", level=1)
hrule()
para("This table is the heart of your pitch. Consider printing it as a one-pager and "
     "putting it on the table during the data-integration session.")

make_table(
    ["Dimension", "DBFZ (resDB / WG Resources)", "PILAR-2b (NIPE-UNICAMP / CP2B)"],
    [
        ["Geography", "Germany & EU", "State of São Paulo (645 municipalities), Brazil"],
        ["Core function", "National biomass-potential monitoring & atlas", "Residue-to-biogas potential platform & SDSS"],
        ["Potential logic", "Theoretical → technical → mobilisable", "Theoretical → mobilisable via FDE correction factors"],
        ["Open access", "Free web app + API (DataLab)", "Live bilingual web app; open-source (GPL-3.0); API"],
        ["Data backbone", "National statistics, residue surveys, time series", "IBGE, MapBiomas, ANP/ANEEL, 50+ residue types (BMP/VS/TS)"],
        ["Spatial stack", "GIS / remote sensing / potential atlas", "PostGIS, React-Leaflet, radius/proximity & infrastructure overlay"],
        ["Decision layer", "Scenarios, system analysis, competing uses", "Technology routes, co-digestion, Leontief IO, payback calculator"],
        ["Maturity", "Established national reference", "Production, INPI-registered, FAPESP-funded, v3.0.3"],
    ],
    widths=[1.4, 2.5, 2.6],
)

callout(
    "The headline message",
    ["Two independently built platforms, same scientific backbone, complementary "
     "geographies and methods. Together they enable a Germany–Brazil comparison of "
     "biomass potentials on harmonised definitions — a publishable, fundable, "
     "policy-relevant goal. PILAR-2b can become the South-America node of a "
     "resDB-style monitoring approach."],
)

page_break()

# ===========================================================================
# 6. COLLABORATION MENU
# ===========================================================================
doc.add_heading("6. Collaboration menu — concrete proposals, ranked", level=1)
hrule()
para("Bring a short menu, not a single demand. Let them co-select; people support "
     "what they help choose. Ordered from lowest-friction (start now) to most "
     "ambitious.")

doc.add_heading("6.1 Quick wins (can start in 2026, no funding needed)", level=2)
numbered("a short, co-authored methods note comparing the DBFZ potential cascade with "
         "the FDE correction-factor approach, using a shared example feedstock "
         "(e.g. sugarcane/vinasse or cattle manure).",
         bold_lead="Joint methods note — ")
numbered("co-present PILAR-2b ↔ resDB interoperability at FOSS4G Europe 2026 "
         "(already on your radar) or a bioenergy venue; a joint abstract is a "
         "low-cost, high-signal commitment.", bold_lead="Joint conference session — ")
numbered("a technical pilot reading the resDB API into PILAR-2b (or vice-versa) to "
         "render a Germany vs São Paulo comparison view — proves interoperability "
         "and gives both sides a demo.", bold_lead="API interoperability pilot — ")

doc.add_heading("6.2 Substantive collaborations (2026–2027)", level=2)
numbered("define a harmonised data schema / ontology for residue types, parameters "
         "(BMP/VS/TS), and the potential cascade, so Brazilian and German figures are "
         "directly comparable.", bold_lead="Harmonised data model — ")
numbered("co-develop the decision-support layer (MCDA — already on your roadmap) for "
         "biogas-plant siting, combining DBFZ’s system-analysis methods with your "
         "spatial proximity/infrastructure engine.", bold_lead="Joint SDSS / MCDA module — ")
numbered("PILAR-2b as a template other Latin-American regions can fork — supporting "
         "DBFZ’s Latin-America strategy with an open, replicable tool.",
         bold_lead="South-America node — ")

doc.add_heading("6.3 Strategic / funded track (2027+)", level=2)
numbered("the 8–12 month internship as the vehicle that delivers the above "
         "(see Section 7).", bold_lead="The post-doctoral internship — ")
numbered("a bilateral research grant (FAPESP–DFG, BMBF Client II, or similar) on "
         "harmonised biomass-potential monitoring & decision support for Brazil and "
         "Germany.", bold_lead="A joint bilateral proposal — ")

callout(
    "How to close Section 6 in the room",
    ["Ask directly: “Of these, which one or two would you find most useful to start "
     "with — and who on your side would be the right person to shape it?” That single "
     "question converts interest into a named owner and a next step."],
)

page_break()

# ===========================================================================
# 7. THE POST-DOC
# ===========================================================================
doc.add_heading("7. The post-doctoral internship — routes, funding, and a draft project",
                level=1)
hrule()

doc.add_heading("7.1 The shape of the ask", level=2)
bullet("8–12 months, in 2027, hosted by WG Resources / Bioenergy Systems Department.",
       bold_lead="Duration & timing — ")
bullet("you bring (or co-apply for) the funding; DBFZ provides supervision, desk, "
       "data access and integration into a live project. This lowers their cost and "
       "raises your odds.", bold_lead="Model — ")
bullet("a working integration between PILAR-2b and DBFZ methods/data + at least one "
       "joint paper — concrete, demonstrable outputs.", bold_lead="Deliverable — ")

doc.add_heading("7.2 Funding routes to research before the trip", level=2)
make_table(
    ["Instrument", "Fit / notes", "Action"],
    [
        ["FAPESP BEPE (Bolsa Estágio de Pesquisa no Exterior)",
         "Natural fit — you are FAPESP-funded (2024/01112-1). BEPE funds research internships abroad attached to your home project.",
         "Confirm eligibility & timing with your PI / FAPESP; align dates to 2027."],
        ["CAPES–Humboldt / Humboldt Fellowship",
         "Prestigious post-doc fellowships for research stays in Germany.",
         "Check eligibility window & deadlines; needs host confirmation."],
        ["DAAD research grants",
         "German agency; supports research stays of varying length for post-docs.",
         "Identify the matching DAAD programme & call dates."],
        ["DFG / BMBF bilateral (FAPESP–DFG, Client II)",
         "Larger, project-level; would fund the broader collaboration, not just you.",
         "Frame as a medium-term goal co-developed with DBFZ."],
        ["DBFZ / project-funded position",
         "If they have a live project needing your skills, a hosted/funded slot may exist.",
         "Ask whether any current project could absorb a visiting researcher."],
    ],
    widths=[1.9, 2.7, 1.9],
)
para("Before the trip, confirm with your PI and FAPESP which instrument is realistic "
     "and what host documents (invitation letter, supervision statement) you would "
     "need. Walking in knowing the funding mechanics signals seriousness.",
     italic=True, color=GREY)

doc.add_heading("7.3 Draft research project (one paragraph to have ready)", level=2)
callout(
    "Working title: “Harmonised, mobilisable biomass-potential assessment and spatial "
    "decision support — a Brazil–Germany comparison”",
    ["Objective: integrate the FDE mobilisable-potential methodology of PILAR-2b with "
     "the DBFZ Resource-Database potential cascade to produce harmonised, comparable "
     "residue-to-biogas potential estimates for São Paulo and a German reference "
     "region, exposed through interoperable open APIs and a shared decision-support "
     "(MCDA) layer for plant siting.",
     "Outputs: (1) a harmonised data model & method paper; (2) a working "
     "PILAR-2b ↔ resDB interoperability demonstrator; (3) a joint MCDA siting module; "
     "(4) a concept note for a follow-on bilateral grant.",
     "Why DBFZ: access to resDB methods/data, system-analysis expertise, and "
     "supervision by a GIS/remote-sensing group with Brazil experience."],
)

page_break()

# ===========================================================================
# 8. AGENDA
# ===========================================================================
doc.add_heading("8. Proposed two-day agenda", level=1)
hrule()
para("Send this (or a trimmed version) ahead of time so hosts can prepare and invite "
     "the right people. Adapt to their constraints — proposing structure is itself a "
     "professional signal.")

doc.add_heading("Day 1 — Mutual understanding", level=2)
make_table(
    ["Block", "Session", "Goal"],
    [
        ["Morning", "Welcome; DBFZ overview & facility tour (labs, DataLab)", "Understand their infrastructure; show genuine interest"],
        ["Late morning", "PILAR-2b live demo (15–20 min) + Q&A", "Establish credibility; trigger ‘we could connect this’ reactions"],
        ["Lunch", "Informal — with WG Resources members", "Personal rapport; learn who does what"],
        ["Afternoon", "Deep-dive: methodology comparison (FDE ↔ DBFZ cascade)", "Find the precise methodological interface"],
        ["End of day", "Recap; capture candidate collaboration ideas", "A shared shortlist to refine on Day 2"],
    ],
    widths=[1.1, 2.9, 2.5],
)

doc.add_heading("Day 2 — From ideas to commitments", level=2)
make_table(
    ["Block", "Session", "Goal"],
    [
        ["Morning", "Data-integration working session (see Section 9)", "Test real interoperability; identify schema gaps"],
        ["Late morning", "Collaboration menu review (Section 6); pick 1–2", "Named owners + next steps for the first joint activity"],
        ["Lunch", "Open / network with department leadership if possible", "Visibility beyond the immediate group"],
        ["Afternoon", "Post-doc discussion: host, timing, funding route", "In-principle agreement + identified funding instrument"],
        ["End of day", "Written wrap-up: who does what, by when", "A short shared action list you both keep"],
    ],
    widths=[1.1, 2.9, 2.5],
)

page_break()

# ===========================================================================
# 9. DATA INTEGRATION
# ===========================================================================
doc.add_heading("9. Data-integration working session — technical preparation", level=1)
hrule()
para("This is where you shine as a developer. Prepare so the session produces "
     "something real, not just slides.")

doc.add_heading("9.1 Bring these, ready to show", level=2)
bullet("a running PILAR-2b instance (live URL + a local/offline fallback in case of "
       "venue Wi-Fi issues);", bold_lead="Live + offline demo — ")
bullet("your API documentation and a couple of example endpoint calls (municipal "
       "potential, proximity analysis);", bold_lead="API docs — ")
bullet("your data dictionary: residue types, parameters (BMP/VS/TS), FDE factors, "
       "units, and source datasets;", bold_lead="Schema / data model — ")
bullet("2–3 sample datasets (e.g. a GeoJSON of municipal potentials, a residue "
       "parameter table) you can hand over or compare on the spot.",
       bold_lead="Sample data — ")

doc.add_heading("9.2 Interoperability checklist to discuss", level=2)
make_table(
    ["Topic", "Questions to resolve together"],
    [
        ["Data formats", "Common exchange formats (GeoJSON, GeoPackage, CSV, Parquet)? Coordinate reference systems & precision?"],
        ["APIs", "Can PILAR-2b consume the resDB API and vice-versa? Auth, rate limits, schema stability?"],
        ["Residue taxonomy", "How do we reconcile 77 (DE) vs 50+ (SP) residue types into a shared ontology/crosswalk?"],
        ["Parameters", "Are BMP/VS/TS definitions and units identical? Reference sources comparable?"],
        ["Potential cascade", "Exact definitions of theoretical / technical / mobilisable — where do FDE factors map?"],
        ["Standards", "OGC services (WMS/WFS/OGC API), INSPIRE alignment, metadata standards?"],
        ["Licensing & open data", "Compatible licences (GPL-3.0, CC-BY)? What can be openly shared vs. restricted?"],
        ["Data sovereignty / GDPR", "Any personal/farm-level data constraints? German/EU data-protection considerations?"],
        ["Versioning & reproducibility", "How are datasets versioned (e.g. Zenodo DOIs as DBFZ uses)? Provenance tracking?"],
    ],
    widths=[1.7, 4.8],
)

callout(
    "Make it tangible",
    ["If venue and time allow, attempt a live pull from the resDB API into a scratch "
     "PILAR-2b view — even a rough Germany-vs-São Paulo bar chart on one feedstock. A "
     "30-minute working prototype is worth more than any slide and demonstrates "
     "exactly the value of hosting you."],
)

page_break()

# ===========================================================================
# 10. QUESTIONS TO ASK THEM
# ===========================================================================
doc.add_heading("10. Questions to ask them", level=1)
hrule()
para("Good questions signal preparation and respect. Group them so you can pick by "
     "context.")

doc.add_heading("Scientific & methodological", level=2)
for q in [
    "How exactly do you define the boundary between technical and mobilisable potential — and how do you handle competing uses and seasonality?",
    "How is the resDB methodology validated against ground truth or plant-level data?",
    "Where do you see the biggest open methodological problems in residue-potential assessment today?",
    "How do you incorporate land-use / remote-sensing change (your background) into potential estimates over time?",
]:
    bullet(q)

doc.add_heading("Collaboration & data", level=2)
for q in [
    "Is the resDB API openly usable for an interoperability pilot with PILAR-2b? Any constraints?",
    "Would a harmonised Brazil–Germany residue taxonomy be useful to your team?",
    "Which current DBFZ projects could most benefit from a São Paulo case study or an open platform like PILAR-2b?",
]:
    bullet(q)

doc.add_heading("Post-doc & logistics", level=2)
for q in [
    "Would your group be open to hosting an 8–12 month visiting post-doc in 2027, and who would supervise?",
    "Which funding instruments have worked best for past visiting researchers from Brazil (DAAD, Humboldt, project funds)?",
    "What working language is used day-to-day in the group, and what onboarding/visa support exists?",
    "What would make a visiting researcher genuinely useful to your team — what should I arrive ready to contribute?",
]:
    bullet(q)

page_break()

# ===========================================================================
# 11. QUESTIONS THEY WILL ASK YOU
# ===========================================================================
doc.add_heading("11. Questions they will ask you — and your answers", level=1)
hrule()
para("Rehearse crisp, honest answers. Anticipating these is half your preparation.")

make_table(
    ["They may ask", "Your prepared line"],
    [
        ["What is PILAR-2b, in one minute?",
         "An open-source geospatial platform that maps mobilisable biogas/bioproduct potential from agricultural, livestock and urban residues across all 645 São Paulo municipalities — from raw data to investment-relevant outputs in the browser."],
        ["What’s genuinely novel about it?",
         "The FDE method: it decomposes theoretical potential into mobilisable potential through explicit, source-attributed factors (collection, competing uses, seasonality, logistics) — transparent and replicable, not a black box."],
        ["How is it different from resDB?",
         "Same backbone, different geography and an explicit factor-based mobilisation step. They’re complementary — which is exactly why integrating them is interesting."],
        ["Who funds and owns it?",
         "FAPESP (2024/01112-1, CP2Bsd), developed at NIPE-UNICAMP, INPI-registered, GPL-3.0 open-source."],
        ["What would you actually do here for 8–12 months?",
         "Deliver a harmonised data model + a working PILAR-2b↔resDB integration + a joint method paper + an MCDA siting module — see the draft project outline."],
        ["What do you need from us?",
         "Supervision, data/method access, and a host letter for the funding application — I aim to bring (or co-apply for) the funding myself."],
        ["What’s your technical level?",
         "Full-stack: FastAPI + PostGIS backend, Next.js/React-Leaflet frontend, APIs, geoprocessing, plus IO economic and Sankey modelling — I build and integrate, not only analyse."],
    ],
    widths=[2.2, 4.3],
)

page_break()

# ===========================================================================
# 12. CHECKLIST
# ===========================================================================
doc.add_heading("12. What to prepare and bring (checklist)", level=1)
hrule()

doc.add_heading("Materials to prepare in advance", level=2)
for item in [
    "A tight 15–20 min PILAR-2b demo (with an offline fallback recording/screenshots).",
    "A one-page collaboration concept note (Appendix A) — printed + PDF.",
    "The PILAR-2b ↔ DBFZ mapping table (Section 5) printed as a leave-behind one-pager.",
    "A 2–3 slide “about me / about PILAR-2b” mini-deck (English).",
    "Updated CV / short academic bio (English) + ORCID / Lattes link.",
    "Draft post-doc research outline (Appendix B), 1 page.",
    "API docs + sample datasets on a USB stick and in cloud storage.",
    "Business cards (with QR to the live platform & repo).",
]:
    bullet(item)

doc.add_heading("Logistics to sort before flying", level=2)
for item in [
    "Confirm meeting dates, location (Leipzig HQ vs. other site), and attendee list.",
    "Confirm the funding instrument you will name (FAPESP BEPE / DAAD / Humboldt) with your PI.",
    "Check passport validity and Schengen entry requirements for your trip dates.",
    "Travel & accommodation in Leipzig; route to DBFZ (Torgauer Straße 116).",
    "Letter/email of introduction from your PI or NIPE, if appropriate.",
    "Test your demo on the actual laptop you will bring; prepare a mobile hotspot.",
]:
    bullet(item)

doc.add_heading("Personal / soft preparation", level=2)
for item in [
    "Rehearse the elevator pitch (Appendix C) until it’s effortless.",
    "Read 2–3 recent DBFZ resDB / WG Resources outputs so you can reference them.",
    "Prepare 3 sincere, specific questions for Dr. Naegeli de Torres about her work.",
    "Learn a few words of German courtesy; have a sentence or two of Portuguese ready in case it builds rapport.",
]:
    bullet(item)

page_break()

# ===========================================================================
# 13. LOGISTICS, LANGUAGE, CULTURE
# ===========================================================================
doc.add_heading("13. Logistics, language and cultural notes", level=1)
hrule()
bullet("DBFZ headquarters is in Leipzig (Torgauer Straße 116). Confirm whether all "
       "sessions are there.", bold_lead="Location — ")
bullet("Working language with international partners is English; written follow-ups in "
       "English are expected and welcome.", bold_lead="Language — ")
bullet("German professional culture values punctuality, structure, directness and "
       "preparation. Arrive early, bring an agenda, follow up promptly and precisely. "
       "Substance over salesmanship.", bold_lead="Etiquette — ")
bullet("Use titles (Dr. / Prof. Dr.) until invited otherwise. A firm handshake and "
       "direct eye contact are standard.", bold_lead="Formality — ")
bullet("Bring more printed leave-behinds than you think you need; Germans often value "
       "a tangible document.", bold_lead="Paper still matters — ")
bullet("June 2026: pleasant weather, but bring smart-casual/business attire for "
       "meetings and comfortable shoes for the facility tour.", bold_lead="Practicalities — ")

# ===========================================================================
# 14. FOLLOW-UP
# ===========================================================================
doc.add_heading("14. Follow-up plan (the 30 days after)", level=1)
hrule()
numbered("send a thank-you email to each person met, with the agreed action list and "
         "any promised materials/data.", bold_lead="Within 48 hours — ")
numbered("deliver the first quick win (e.g. share the harmonised-schema draft or the "
         "API pilot repo) to prove momentum.", bold_lead="Within 2 weeks — ")
numbered("circulate a short joint outline (method note or conference abstract) and "
         "propose the post-doc funding-application timeline.", bold_lead="Within 1 month — ")
numbered("keep a shared document of decisions and next steps; quiet collaborations "
         "die — visible momentum gets funded.", bold_lead="Ongoing — ")

callout(
    "The principle behind the whole trip",
    ["You are not asking DBFZ for a favour. You are offering a peer institution a "
     "ready-made platform, a novel method, a new geography, and a motivated builder — "
     "and proposing that the post-doc is simply the most efficient way to capture that "
     "shared value. Lead with the collaboration; the internship follows."],
)

page_break()

# ===========================================================================
# APPENDIX A — CONCEPT NOTE
# ===========================================================================
doc.add_heading("Appendix A — One-page collaboration concept note (draft)", level=1)
hrule()
para("Draft text you can adapt into a printed leave-behind or an email.", italic=True,
     color=GREY)

para("UNICAMP (NIPE / CP2B) — DBFZ collaboration concept", bold=True, color=DARK,
     size=12)
rich([("Context. ", {"bold": True}),
      ("The DBFZ Resource Database provides Germany’s national, open monitoring of "
       "biomass potentials. PILAR-2b, developed at NIPE-UNICAMP and funded by FAPESP, "
       "provides an equivalent open platform for the State of São Paulo, adding an "
       "explicit feedstock-dependent efficiency (FDE) method to estimate mobilisable "
       "potential. The two systems share a scientific backbone and complementary "
       "geographies.", {})])
rich([("Proposal. ", {"bold": True}),
      ("Establish a UNICAMP–DBFZ collaboration to (1) harmonise residue taxonomies and "
       "potential definitions; (2) demonstrate PILAR-2b ↔ resDB data/API "
       "interoperability; (3) co-develop a spatial decision-support (MCDA) layer for "
       "biogas-plant siting; and (4) co-author a method paper comparing the DBFZ "
       "potential cascade with the FDE approach.", {})])
rich([("Vehicle. ", {"bold": True}),
      ("An 8–12 month post-doctoral research internship at DBFZ in 2027 (funding to be "
       "secured via FAPESP BEPE / DAAD / Humboldt), delivering the integration and the "
       "joint paper, and seeding a follow-on bilateral grant proposal.", {})])
rich([("Mutual benefit. ", {"bold": True}),
      ("DBFZ extends its Latin-America strategy to São Paulo with a ready open platform "
       "and a builder; UNICAMP gains methodological alignment with a world reference "
       "and an international post-doc placement.", {})])
rich([("Contact. ", {"bold": True}),
      ("Lucas Nakamura Cerejo — NIPE-UNICAMP / CP2B — lucasnc@unicamp.br — "
       "live platform: cp2b.unicamp.br/pilar2b — code: github.com/aikiesan/Pilar-2b",
       {})])

page_break()

# ===========================================================================
# APPENDIX B — RESEARCH OUTLINE
# ===========================================================================
doc.add_heading("Appendix B — Draft post-doc research outline", level=1)
hrule()

para("Title (working)", bold=True, color=DARK)
para("Harmonised, mobilisable biomass-potential assessment and spatial decision "
     "support: a Brazil–Germany comparison.")

para("Problem", bold=True, color=DARK)
para("Biomass-potential estimates are hard to compare across countries because "
     "residue taxonomies, parameters and the theoretical→mobilisable step differ. This "
     "blocks cross-border learning and policy transfer.")

para("Objectives", bold=True, color=DARK)
numbered("harmonise residue taxonomies and BMP/VS/TS parameter definitions between "
         "resDB and PILAR-2b;")
numbered("formally map the FDE correction factors onto the DBFZ "
         "theoretical→technical→mobilisable cascade;")
numbered("build a working PILAR-2b ↔ resDB interoperability demonstrator (shared "
         "APIs / data model);")
numbered("co-develop an MCDA spatial decision-support layer for biogas-plant siting;")
numbered("produce a harmonised Brazil (São Paulo) vs Germany comparison on selected "
         "feedstocks (e.g. sugarcane/vinasse, cattle manure).")

para("Methods", bold=True, color=DARK)
para("GIS & geoprocessing (PostGIS, GeoPandas), API integration, potential-cascade "
     "modelling, multi-criteria decision analysis, open-data publication (Zenodo "
     "DOIs).")

para("Deliverables", bold=True, color=DARK)
bullet("a harmonised data-model specification & crosswalk;")
bullet("a peer-reviewed method paper (FDE ↔ DBFZ cascade);")
bullet("an open interoperability demonstrator + documentation;")
bullet("an MCDA siting module integrated in PILAR-2b;")
bullet("a concept note for a follow-on FAPESP–DFG / BMBF bilateral grant.")

para("Timeline (8–12 months, 2027)", bold=True, color=DARK)
make_table(
    ["Phase", "Months", "Focus"],
    [
        ["Onboarding & alignment", "1–2", "Method deep-dive; data access; finalise scope"],
        ["Harmonisation", "3–5", "Taxonomy crosswalk; parameter & cascade mapping"],
        ["Integration", "5–8", "API interoperability demonstrator; comparison runs"],
        ["Decision support", "8–10", "MCDA siting module; case study"],
        ["Writing & next grant", "10–12", "Method paper; bilateral-grant concept note"],
    ],
    widths=[2.2, 1.0, 3.3],
)

page_break()

# ===========================================================================
# APPENDIX C — PITCH
# ===========================================================================
doc.add_heading("Appendix C — Elevator pitch & key talking points", level=1)
hrule()

para("30-second pitch", bold=True, color=DARK)
para("“I’m Lucas, from NIPE-UNICAMP in Brazil. I built PILAR-2b — an open-source "
     "platform that maps mobilisable biogas potential from residues across all 645 "
     "municipalities of São Paulo. It’s essentially the São Paulo counterpart of your "
     "Resource Database, with an added method — the FDE — for getting from theoretical "
     "to mobilisable potential. I’m here because integrating the two could give us a "
     "harmonised Brazil–Germany picture, and I’d love to explore doing that with your "
     "group.”", italic=True)

para("Three messages to repeat", bold=True, color=DARK)
numbered("Same backbone, complementary geographies — PILAR-2b and resDB are natural "
         "partners, not competitors.")
numbered("The FDE method is a real, transparent contribution to the "
         "theoretical→mobilisable step.")
numbered("I build and integrate — a hosted post-doc would produce working software and "
         "a joint paper, not just analysis.")

para("Three things to avoid", bold=True, color=DARK)
bullet("Leading with the post-doc request before establishing mutual value.")
bullet("Overselling — be precise about what PILAR-2b does and does not yet do "
       "(e.g. MCDA is on the roadmap).")
bullet("Vague asks — always end a topic with a concrete next step and an owner.")

page_break()

# ===========================================================================
# APPENDIX D — SOURCES
# ===========================================================================
doc.add_heading("Appendix D — Sources & useful links", level=1)
hrule()
para("Key references consulted in preparing this dossier (verify details live before "
     "the trip):", italic=True, color=GREY)

links = [
    ("DBFZ — Bioenergy Systems Department",
     "https://www.dbfz.de/en/the-dbfz/research-departments/bioenergy-systems-department"),
    ("DBFZ — Working Group Resources",
     "https://www.dbfz.de/en/working-group-resources"),
    ("Dr. Friederike Naegeli de Torres — DBFZ staff profile",
     "https://www.dbfz.de/das-dbfz/personalverzeichnis/employee/friederike-naegeli-de-torres/page/detail"),
    ("Dr. Friederike Naegeli de Torres — LinkedIn",
     "https://www.linkedin.com/in/fnaegelidetorres/"),
    ("DBFZ Resource Database (resDB) — DataLab",
     "https://datalab.dbfz.de/resdb"),
    ("DBFZ DataLab — OpenData",
     "https://datalab.dbfz.de/home/"),
    ("DBFZ — Country Profiles: Brazil",
     "https://www.dbfz.de/en/research-departments/international/country-profiles-brazil.html"),
    ("DBFZ — International activities",
     "https://www.dbfz.de/en/research/international-activities/"),
    ("DBFZ — Research & Development",
     "https://www.dbfz.de/en/research/research-and-development/"),
    ("DBFZ Resource Database — DE-Biomass Monitor (Zenodo)",
     "https://zenodo.org/records/14273285"),
    ("PILAR-2b — live platform",
     "https://cp2b.unicamp.br/pilar2b/pt-BR"),
    ("PILAR-2b — source code",
     "https://github.com/aikiesan/Pilar-2b"),
    ("NIPE-UNICAMP / CP2B",
     "https://nipe.unicamp.br/cp2b"),
]
for title, url in links:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(title + " — ")
    r.font.bold = True
    r.font.size = Pt(10)
    ru = p.add_run(url)
    ru.font.size = Pt(9.5)
    ru.font.color.rgb = GREEN

para("", space_after=6)
para("Note: a few DBFZ pages block automated retrieval; confirm the latest figures, "
     "project names and Dr. Naegeli de Torres’s current title directly on the DBFZ "
     "site and LinkedIn shortly before you travel.", italic=True, color=GREY, size=9)

# ---------------------------------------------------------------------------
# Footer with page numbers
# ---------------------------------------------------------------------------
section = doc.sections[0]
footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = fp.add_run("DBFZ Trip Preparation Dossier · Lucas N. Cerejo · NIPE-UNICAMP / CP2B · ")
run.font.size = Pt(8)
run.font.color.rgb = GREY
# page number field
fldSimple = OxmlElement("w:fldSimple")
fldSimple.set(qn("w:instr"), "PAGE")
fp._p.append(fldSimple)

out = "docs/DBFZ_Germany_Trip_Preparation.docx"
doc.save(out)
print("Saved:", out)
