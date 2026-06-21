#!/usr/bin/env python3
"""
Generator for the DBFZ "People & Relationships" dossier (.docx).

Companion to DBFZ_Germany_Trip_Preparation.docx. Deep profiles of the key people:
Dr. Friederike Naegeli de Torres (scientific host), Susann Günther (her peer /
methodological twin), Dr. Sven Schaller (MOU partner / warm anchor), and how
Prof. Bruna Moraes (CP2B leader) and the NIPE-DBFZ MOU tie it together.

Run: python3 docs/build_dbfz_people_doc.py
Output: docs/DBFZ_People_and_Collaboration_Profiles.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

GREEN = RGBColor(0x2E, 0x7D, 0x32)
DARK = RGBColor(0x1B, 0x2A, 0x3A)
GREY = RGBColor(0x55, 0x5F, 0x6B)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
LIGHT_FILL = "EAF3EC"
ALT_FILL = "F4F8F5"
BLUE_FILL = "E7EEF6"
AMBER_FILL = "FFF4E5"

doc = Document()

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.12

for name, size, color in [("Heading 1", 17, GREEN), ("Heading 2", 13.5, DARK),
                          ("Heading 3", 11.5, BLUE)]:
    st = doc.styles[name]
    st.font.name = "Calibri"
    st.font.size = Pt(size)
    st.font.color.rgb = color
    st.font.bold = True
    st.paragraph_format.space_before = Pt(12)
    st.paragraph_format.space_after = Pt(4)


def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), fill)
    tcPr.append(sh)


def set_cell_text(cell, text, bold=False, color=None, size=10, align=None, fill=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = color
    if fill:
        shade(cell, fill)
    return cell


def para(text="", bold=False, italic=False, color=None, size=None, align=None,
         space_after=None, space_before=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    if text:
        r = p.add_run(text)
        r.font.bold = bold
        r.font.italic = italic
        if color:
            r.font.color.rgb = color
        if size:
            r.font.size = Pt(size)
    return p


def rich(parts, align=None, space_after=None, style=None):
    p = doc.add_paragraph(style=style)
    if align:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    for text, fmt in parts:
        r = p.add_run(text)
        r.font.bold = fmt.get("bold", False)
        r.font.italic = fmt.get("italic", False)
        if fmt.get("color"):
            r.font.color.rgb = fmt["color"]
        if fmt.get("size"):
            r.font.size = Pt(fmt["size"])
    return p


def bullet(text, bold_lead=None, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    if bold_lead:
        r = p.add_run(bold_lead)
        r.font.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def numbered(text, bold_lead=None):
    p = doc.add_paragraph(style="List Number")
    if bold_lead:
        r = p.add_run(bold_lead)
        r.font.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def hrule(color="2E7D32"):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    pPr.append(pbdr)
    p.paragraph_format.space_after = Pt(2)
    return p


def make_table(headers, rows, widths=None, header_fill=LIGHT_FILL, zebra=True,
               font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, color=DARK, size=font_size, fill=header_fill)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci, val in enumerate(row):
            set_cell_text(cells[ci], val, size=font_size)
            if zebra and ri % 2 == 1:
                shade(cells[ci], ALT_FILL)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    return t


def profile_card(rows, fill=BLUE_FILL):
    t = doc.add_table(rows=0, cols=2)
    t.style = "Table Grid"
    for k, v in rows:
        cells = t.add_row().cells
        set_cell_text(cells[0], k, bold=True, color=DARK, size=9.5, fill=fill)
        set_cell_text(cells[1], v, size=9.5)
        cells[0].width = Inches(1.6)
        cells[1].width = Inches(4.9)
    return t


def callout(title, body_lines, fill=LIGHT_FILL, title_color=GREEN):
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    cell = t.rows[0].cells[0]
    shade(cell, fill)
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.font.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = title_color
    for line in body_lines:
        bp = cell.add_paragraph()
        bp.paragraph_format.space_before = Pt(2)
        rr = bp.add_run(line)
        rr.font.size = Pt(10)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def page_break():
    doc.add_page_break()


# ===========================================================================
# COVER
# ===========================================================================
doc.add_paragraph().paragraph_format.space_before = Pt(50)
para("PEOPLE & COLLABORATION DOSSIER", bold=True, color=GREEN, size=12,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para("Who's Who at DBFZ — Profiles, Research Interests & How to Engage",
     bold=True, color=DARK, size=22, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para("Companion to the DBFZ Germany Trip Preparation Dossier",
     italic=True, color=GREY, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
para("Built around the existing NIPE–DBFZ relationship and the MOU in progress",
     color=GREEN, bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)

meta = doc.add_table(rows=0, cols=2)
meta.alignment = WD_TABLE_ALIGNMENT.CENTER
for k, v in [
    ("Prepared for", "Lucas Nakamura Cerejo — NIPE-UNICAMP / CP2B"),
    ("Profiles inside", "Dr. Friederike Naegeli de Torres · Susann Günther · "
     "Dr. Sven Schaller · Prof. Bruna Moraes (CP2B)"),
    ("Existing bridge", "MOU in progress with Dr. Sven Schaller (DBFZ) ↔ "
     "Prof. Bruna Moraes (CP2B leader)"),
    ("Your target host", "Working Group “Resources”, Bioenergy Systems Dept., DBFZ"),
    ("Document date", "21 June 2026"),
]:
    cells = meta.add_row().cells
    set_cell_text(cells[0], k, bold=True, color=DARK, size=10, fill=LIGHT_FILL)
    set_cell_text(cells[1], v, size=10)
    cells[0].width = Inches(1.7)
    cells[1].width = Inches(4.6)

para("", space_after=14)
para("Note on names/titles: confirm exact spellings and current titles on the DBFZ "
     "staff directory before the trip. This dossier uses “Susann Günther” (DBFZ "
     "directory spelling) and treats her academic title as to-be-confirmed.",
     italic=True, color=GREY, size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER)

page_break()

# ===========================================================================
# 0. HOW THE PEOPLE FIT TOGETHER
# ===========================================================================
doc.add_heading("1. The relationship map — how the people fit together", level=1)
hrule()
para("You are not arriving cold. There is already a warm institutional channel "
     "between NIPE/CP2B and DBFZ. Your job is to travel along that channel into the "
     "scientific group that matches your profile, and to make the post-doc the "
     "natural next step under the MOU.")

make_table(
    ["Person", "Side", "Role in your strategy"],
    [
        ["Prof. Bruna Moraes", "NIPE / CP2B (Brazil)",
         "Your leader & champion. Leads CP2B; holds the relationship with Schaller; sponsors the MOU and your mobility."],
        ["Dr. Sven Schaller", "DBFZ (Germany)",
         "Your warm anchor & door-opener. International Knowledge & Technology Transfer; runs the Brazil network; your MOU counterpart."],
        ["Dr. Friederike Naegeli de Torres", "DBFZ (Germany)",
         "Your scientific host. Head of WG “Resources”; GIS/biomass-potential — the best match for PILAR-2b and a post-doc supervisor."],
        ["Susann Günther", "DBFZ (Germany)",
         "Your methodological twin / peer. Spatial mapping of theoretical biomass potential across Europe — closest to your day-to-day work."],
        ["You — Lucas N. Cerejo", "NIPE / CP2B (Brazil)",
         "The builder. PILAR-2b + FDE method; the person who turns the MOU into working software and joint papers."],
    ],
    widths=[1.9, 1.5, 3.1],
)

callout(
    "The path in one sentence",
    ["Bruna + Schaller + the MOU give you the warm institutional bridge; Friederike + "
     "Günther give you the scientific home; PILAR-2b + the FDE method give you the "
     "deliverable — and the 8–12 month post-doc is simply the most efficient way to "
     "connect all three."],
)

page_break()

# ===========================================================================
# 2. SVEN SCHALLER
# ===========================================================================
doc.add_heading("2. Dr. Sven Schaller — your warm anchor & MOU partner", level=1)
hrule()
para("Already friendly with NIPE; the relationship is good and the MOU is nearly "
     "done. Treat him as your sponsor inside DBFZ — the person who introduces you to "
     "the right scientists and frames your visit within the institutional agreement.")

profile_card([
    ("Role at DBFZ", "Coordinator for International Knowledge- & Technology Transfer"),
    ("What he runs", "DBFZ's international cooperation, especially the Brazil network (BiReSb) and partner agreements"),
    ("Brazil track record", "DBFZ active in Brazil since 2005; ~15 cooperation agreements; regional focus on South & South-East Brazil"),
    ("Thematic focus", "Biogas from agricultural & municipal residues; sustainable biogas from bioethanol residues (vinasse) and manure; pilot plants"),
    ("Relationship status", "Warm — strong existing relationship with NIPE; MOU with CP2B (Prof. Bruna Moraes) close to signature"),
    ("Contact", "+49 (0)341 2434-551 · DBFZ staff directory"),
], fill=AMBER_FILL)

doc.add_heading("How to use the relationship with Schaller", level=2)
bullet("publicly thank him and CP2B/Bruna for the bridge; acknowledge the MOU as the "
       "frame for everything you propose.", bold_lead="Anchor to the MOU — ")
bullet("ask him to introduce you to Friederike Naegeli de Torres and Susann Günther, "
       "and to vouch for the post-doc idea internally.", bold_lead="Ask for introductions — ")
bullet("frame your São Paulo/PILAR-2b work as extending DBFZ's Brazil network from the "
       "South into the largest sugarcane/bioenergy state — directly serving his "
       "mandate.", bold_lead="Serve his agenda — ")
bullet("position the 8–12 month internship as a concrete deliverable under the MOU — "
       "the kind of mobility that makes the agreement live, not just paper.",
       bold_lead="Make the MOU tangible — ")

callout(
    "Talking point for Schaller",
    ["“The MOU between CP2B and DBFZ gives us the frame; PILAR-2b gives us a ready "
     "platform for São Paulo. The most efficient way to turn the MOU into real output "
     "would be a post-doc placement where I integrate our São Paulo data with DBFZ "
     "methods. Could you help connect me with Friederike's Resources group to shape "
     "that?”"],
    fill=AMBER_FILL, title_color=RGBColor(0xB0, 0x6A, 0x00),
)

page_break()

# ===========================================================================
# 3. FRIEDERIKE — DEEP DIVE
# ===========================================================================
doc.add_heading("3. Dr. Friederike Naegeli de Torres — your scientific host (deep dive)",
                level=1)
hrule()
para("This is the most important profile in the dossier. She leads the group you want "
     "to join and is the most likely supervisor for your post-doc. Know her work well "
     "enough to discuss it specifically.")

profile_card([
    ("Role", "Head of the Working Group “Resources”, DBFZ (since 2023); researcher & GIS expert at DBFZ since 2020"),
    ("Training", "Geographer · MSc in Technology & Resource Management in the Tropics & Subtropics · PhD in Remote Sensing & Geoinformatics"),
    ("Guiding question", "“Where can we find how much of which kind of biomass? How much is mobilisable, and what are the utilisation paths?”"),
    ("Core expertise", "Spatial assessment of biomass potentials (national & international); GIS; remote sensing; resource monitoring; biogenic-waste utilisation"),
    ("Geographies", "Brazil, Colombia, Togo, Ethiopia — strong tropics/subtropics & Global-South orientation"),
    ("Contact", "+49 (0)341 2434-620 · DBFZ staff directory"),
])

doc.add_heading("3.1 Research topics & interests — and how PILAR-2b connects to each", level=2)
make_table(
    ["Her topic / interest", "Evidence", "Your hook (what to say / show)"],
    [
        ["Spatial assessment of biomass potential",
         "Core of WG Resources; the resDB potential atlas",
         "PILAR-2b is exactly this for São Paulo's 645 municipalities — show the choropleth live."],
        ["Theoretical → mobilisable potential",
         "DBFZ potential cascade; resDB methodology",
         "Your FDE method makes the mobilisation step explicit & factor-attributed — a methodological contribution."],
        ["Resource monitoring & databases",
         "DBFZ Resource Database / DE-Biomass Monitor (time series)",
         "PILAR-2b as a São Paulo node; propose API interoperability and a harmonised data model."],
        ["GIS & remote sensing",
         "PhD in remote sensing & geoinformatics",
         "Your PostGIS/Leaflet stack + MapBiomas land-use integration; speak her technical language."],
        ["International / Global-South biomass mapping",
         "LabTogo project; Colombia, Ethiopia work",
         "PILAR-2b is a replicable open template other Latin-American regions can fork."],
        ["Brazil — land use, carbon, pastures",
         "Atlantic-Forest carbon study (Rio de Janeiro); INTECRAL pasture rehabilitation",
         "Connect residue mobilisation to land-use & sustainability; you share the Brazilian context."],
        ["Biogenic waste & residue utilisation paths",
         "WG Resources scope",
         "Your technology-routes + co-digestion + IO economic modules answer the 'utilisation paths' question."],
    ],
    widths=[1.9, 1.9, 2.7],
)

doc.add_heading("3.2 Notable projects & outputs to be able to reference", level=2)
bullet("national biomass monitoring published as an open web app + API; DE-Biomass "
       "Monitor with 2010–2020 time series. The reference point for any "
       "interoperability discussion.", bold_lead="DBFZ Resource Database (resDB) — ")
bullet("biomass potential analysis and set-up of research capacities for a biogas "
       "sector in Togo — a model for capacity-building you could echo for "
       "Latin America.", bold_lead="LabTogo project — ")
bullet("temporal & spatial mapping of theoretical biomass potential across the EU "
       "(also Günther's theme) — the EU-scale analogue of your São Paulo mapping.",
       bold_lead="EU theoretical biomass potential mapping — ")
bullet("carbon-storage capacity of a fragmented Atlantic-Forest landscape (Rio de "
       "Janeiro) and pastures as climate-mitigation opportunity areas (INTECRAL). "
       "Her Brazilian and land-use/carbon credentials.",
       bold_lead="Brazil land-use & carbon work — ")

doc.add_heading("3.3 How to engage her — conversation hooks", level=2)
bullet("open with genuine interest in her work, not your ask. Ask how she defines the "
       "boundary between technical and mobilisable potential — then show how FDE "
       "handles it.", bold_lead="Lead with her science — ")
bullet("a few words in Portuguese, or referencing the Atlantic Forest / her INTECRAL "
       "work, signals you've done your homework and share her context.",
       bold_lead="Brazil connection — ")
bullet("frame PILAR-2b as a peer system to resDB, then ask where she sees the most "
       "useful point of integration.", bold_lead="Peer, not petitioner — ")
bullet("ask what would make a visiting researcher genuinely useful to her group — let "
       "her describe the ideal post-doc, then map yourself onto it.",
       bold_lead="Co-design the role — ")

callout(
    "Three specific questions for Dr. Naegeli de Torres",
    ["• “How do you currently handle competing uses and seasonality when moving from "
     "technical to mobilisable potential — and would an explicit factor-based approach "
     "like our FDE be useful to compare against?”",
     "• “Given your LabTogo and Latin-America experience, would a replicable open "
     "platform for São Paulo residues be valuable to your group's international work?”",
     "• “If we hosted a Brazil–Germany comparison on a few feedstocks (vinasse, "
     "manure), what would make it methodologically robust in your view?”"],
)

page_break()

# ===========================================================================
# 4. SUSANN GÜNTHER
# ===========================================================================
doc.add_heading("4. Susann Günther — your peer & methodological twin", level=1)
hrule()
para("If Friederike is your potential supervisor, Günther is likely your closest "
     "day-to-day collaborator. Her published work overlaps almost exactly with what "
     "you do — which makes her both the easiest person to talk shop with and a natural "
     "co-author.")

profile_card([
    ("Role", "Researcher, DBFZ Bioenergy Systems Department (Bioenergiesysteme)"),
    ("Signature work", "“Temporal and Spatial Mapping of the Theoretical Biomass Potential of 11 Residues Across Europe”"),
    ("Focus areas", "Feedstock characterisation, localisation & potential estimation; supply-chain & feedstock-sustainability assessment; biogas systems"),
    ("Platform link", "Contributor to the DBFZ Resource Database / DE-Biomass Monitor"),
    ("Networks", "DBFZ's IEA Bioenergy involvement; EUBCE conference community"),
    ("Contact", "DBFZ staff/employee directory (Bioenergiesysteme)"),
])

doc.add_heading("4.1 Why she matters to you", level=2)
bullet("her EU-wide spatial mapping of theoretical biomass potential is the European "
       "mirror of your São Paulo mapping — same methods, different geography.",
       bold_lead="Direct overlap — ")
bullet("feedstock characterisation, localisation and potential estimation are exactly "
       "the columns in your residue database (BMP/VS/TS, 50+ residue types).",
       bold_lead="Shared data model — ")
bullet("a co-authored Brazil–Germany comparison paper is a low-friction first "
       "deliverable that plays to both your strengths.", bold_lead="Natural co-author — ")
bullet("she sits in the same department flow as the Resources group, so befriending "
       "her strengthens your case to be hosted.", bold_lead="Internal ally — ")

doc.add_heading("4.2 How to engage her — conversation hooks", level=2)
bullet("ask how she selected and harmonised the 11 residues across Europe, and how she "
       "handled data gaps between countries — then offer the São Paulo residue set as "
       "a comparison case.", bold_lead="Talk methods — ")
bullet("propose a concrete mini-collaboration: align a shared residue taxonomy and "
       "compare theoretical potentials for 2–3 common feedstocks.",
       bold_lead="Offer a first joint task — ")
bullet("explore co-submitting to EUBCE or a bioenergy journal — she knows those "
       "venues.", bold_lead="Shared venues — ")

callout(
    "Talking point for Günther",
    ["“Your EU mapping of theoretical biomass potential is essentially what we did for "
     "São Paulo — so a Brazil–Germany comparison on a shared residue taxonomy almost "
     "writes itself. Could we pick two or three common feedstocks and try a harmonised "
     "comparison together?”"],
)
para("Small but important: confirm her academic title before addressing her. Public "
     "records list her with an M.Sc.; use “Ms. Günther” or her full name until you "
     "know whether she holds a doctorate.", italic=True, color=GREY, size=9)

page_break()

# ===========================================================================
# 5. BRUNA MORAES
# ===========================================================================
doc.add_heading("5. Prof. Bruna Moraes (CP2B leader) — your champion at home", level=1)
hrule()
para("She leads CP2B, is co-author on PILAR-2b, holds the relationship with Schaller, "
     "and sponsors the MOU. Your mobility and funding case run through her — align "
     "tightly before and during the trip.")

profile_card([
    ("Role", "Leader of CP2B (NIPE-UNICAMP); co-author / PI context for PILAR-2b"),
    ("Relationship to DBFZ", "Holds the warm relationship with Dr. Sven Schaller; sponsor of the NIPE–DBFZ MOU"),
    ("Your dependency on her", "Endorsement for the post-doc, alignment with the MOU scope, and support for the funding route (e.g. FAPESP BEPE)"),
], fill=LIGHT_FILL)

doc.add_heading("How to leverage this before & during the trip", level=2)
bullet("agree with Bruna exactly what you may propose at DBFZ (collaboration scope, "
       "the post-doc, the funding instrument) so you speak with one institutional "
       "voice.", bold_lead="Pre-align — ")
bullet("reference the MOU and Bruna–Schaller relationship when you open with the DBFZ "
       "team — it signals legitimacy and continuity.", bold_lead="Cite the bridge — ")
bullet("ask Bruna whether the post-doc can be framed as a deliverable of the MOU, "
       "which strengthens both the MOU and your application.",
       bold_lead="Tie post-doc to MOU — ")
bullet("after the trip, report back to Bruna with the agreed next steps so she can "
       "reinforce them with Schaller at the institutional level.",
       bold_lead="Close the loop — ")

page_break()

# ===========================================================================
# 6. PUTTING IT TOGETHER
# ===========================================================================
doc.add_heading("6. Putting it together — your engagement sequence", level=1)
hrule()

doc.add_heading("6.1 The order of engagement", level=2)
numbered("thank Schaller and acknowledge the MOU/Bruna bridge; ask him to connect you "
         "into the Resources group.", bold_lead="Start warm (Schaller) — ")
numbered("lead with her science, show PILAR-2b as a peer to resDB, and co-design what "
         "a useful post-doc would look like.", bold_lead="Win the host (Friederike) — ")
numbered("talk methods, propose a concrete comparison, and float a co-authored paper.",
         bold_lead="Recruit the peer (Günther) — ")
numbered("keep Bruna aligned throughout; frame the post-doc as an MOU deliverable.",
         bold_lead="Anchor at home (Bruna) — ")

doc.add_heading("6.2 One-line message tuned to each person", level=2)
make_table(
    ["Person", "Your one-line message"],
    [
        ["Dr. Sven Schaller",
         "“Let's make the MOU tangible — a São Paulo post-doc that turns PILAR-2b into joint DBFZ output.”"],
        ["Dr. Friederike Naegeli de Torres",
         "“PILAR-2b is the São Paulo counterpart of your Resource Database; let's integrate them and compare potentials.”"],
        ["Susann Günther",
         "“Your EU residue mapping and our São Paulo mapping are twins — let's co-author a Brazil–Germany comparison.”"],
        ["Prof. Bruna Moraes",
         "“The post-doc can be the first concrete deliverable of the DBFZ MOU — here's the plan and funding route.”"],
    ],
    widths=[2.2, 4.3],
)

callout(
    "The principle",
    ["Every person here already has a reason to say yes: Schaller wants the Brazil "
     "network to grow; Friederike wants strong international resource-mapping; Günther "
     "wants comparison data and co-authors; Bruna wants the MOU to produce results. "
     "Your job is simply to show each of them that you are the most efficient way to "
     "get what they already want."],
)

doc.add_heading("6.3 Pre-trip verification checklist", level=2)
for item in [
    "Confirm exact spellings & current titles (Naegeli de Torres, Günther, Schaller) on the DBFZ staff directory.",
    "Confirm with Bruna the agreed scope you may propose and the funding instrument to name.",
    "Read one recent output from Friederike's group and one of Günther's mapping abstracts so you can cite them specifically.",
    "Confirm whether Günther / others should be invited to your demo and the data-integration session.",
    "Check the MOU's status and scope so you can reference it accurately in the room.",
]:
    bullet(item)

page_break()

# ===========================================================================
# 7. SOURCES
# ===========================================================================
doc.add_heading("7. Sources & useful links", level=1)
hrule()
para("Confirm details live before the trip — some DBFZ pages block automated "
     "retrieval and titles/projects change.", italic=True, color=GREY)

links = [
    ("Dr. Friederike Naegeli de Torres — DBFZ profile",
     "https://www.dbfz.de/das-dbfz/personalverzeichnis/employee/friederike-naegeli-de-torres/page/detail"),
    ("Dr. Friederike Naegeli de Torres — LinkedIn",
     "https://www.linkedin.com/in/fnaegelidetorres/"),
    ("DBFZ — Working Group Resources",
     "https://www.dbfz.de/en/working-group-resources"),
    ("Susann Günther — DBFZ contact",
     "https://www.dbfz.de/en/the-dbfz/employee-directory/employee/susann-guenther/page/kontakt"),
    ("Susann Günther — LinkedIn",
     "https://www.linkedin.com/in/susann-g%C3%BCnther-a842b256/"),
    ("Susann Günther — ResearchGate profile",
     "https://www.researchgate.net/profile/Guenther-Susann"),
    ("Dr. Sven Schaller — DBFZ contact",
     "https://www.dbfz.de/en/the-dbfz/employee-directory/employee/sven-schaller/page/kontakt"),
    ("DBFZ — Country Profiles: Brazil",
     "https://www.dbfz.de/en/research-departments/international/country-profiles-brazil.html"),
    ("DBFZ — International activities (America)",
     "https://www.dbfz.de/en/research/international-activities/america"),
    ("DBFZ Resource Database (resDB) — DataLab",
     "https://datalab.dbfz.de/resdb"),
    ("LabTogo project — biomass potential & biogas capacity",
     "https://www.tib-op.org/ojs/index.php/th-wildau-ensp/article/view/14"),
    ("DBFZ — Bioenergy Systems Department",
     "https://www.dbfz.de/en/the-dbfz/research-departments/bioenergy-systems-department"),
]
for title, url in links:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(title + " — ")
    r.font.bold = True
    r.font.size = Pt(10)
    ru = p.add_run(url)
    ru.font.size = Pt(9.5)
    ru.font.color.rgb = GREEN

# footer
section = doc.sections[0]
fp = section.footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = fp.add_run("DBFZ People & Collaboration Dossier · Lucas N. Cerejo · NIPE-UNICAMP / CP2B · ")
run.font.size = Pt(8)
run.font.color.rgb = GREY
fld = OxmlElement("w:fldSimple")
fld.set(qn("w:instr"), "PAGE")
fp._p.append(fld)

out = "docs/DBFZ_People_and_Collaboration_Profiles.docx"
doc.save(out)
print("Saved:", out)
